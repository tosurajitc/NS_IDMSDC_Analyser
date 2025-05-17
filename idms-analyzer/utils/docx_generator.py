"""
docx_generator.py - Handles conversion of content to Word document format.
This module provides functionality to create well-formatted Word documents
from markdown content with proper styling and structure.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import re
from io import BytesIO
from datetime import datetime
import base64
import logging
from typing import Optional, Dict, Any, List, Tuple

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocxGenerator:
    """Class for generating well-formatted Word documents"""
    
    def __init__(self):
        """Initialize the DocxGenerator"""
        self.doc = None
    
    def create_document(self, title: str) -> Document:
        """
        Create a new Word document with basic styling.
        
        Args:
            title: Document title
            
        Returns:
            Document object
        """
        self.doc = Document()
        
        # Set up styles
        self._setup_document_styles()
        
        # Add title page
        self._add_title_page(title)
        
        # Add page break after title
        self.doc.add_page_break()
        
        return self.doc
    
    def _setup_document_styles(self) -> None:
        """Set up document styles for consistent formatting"""
        # Modify the Normal style
        style_normal = self.doc.styles['Normal']
        font = style_normal.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Modify Heading styles
        for i in range(1, 5):
            style_name = f'Heading {i}'
            if style_name in self.doc.styles:
                style = self.doc.styles[style_name]
                font = style.font
                font.name = 'Calibri'
                
                # Different sizes for different heading levels
                if i == 1:
                    font.size = Pt(16)
                    font.bold = True
                    font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
                elif i == 2:
                    font.size = Pt(14)
                    font.bold = True
                    font.color.rgb = RGBColor(0, 51, 102)
                elif i == 3:
                    font.size = Pt(12)
                    font.bold = True
                    font.color.rgb = RGBColor(0, 51, 102)
                else:
                    font.size = Pt(11)
                    font.bold = True
                    font.italic = True
        
        # Create a Code style
        if 'Code' not in self.doc.styles:
            code_style = self.doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            font = code_style.font
            font.name = 'Courier New'
            font.size = Pt(9)
            paragraph_format = code_style.paragraph_format
            paragraph_format.space_before = Pt(6)
            paragraph_format.space_after = Pt(6)
            paragraph_format.left_indent = Inches(0.5)
    
    def _add_title_page(self, title: str) -> None:
        """
        Add a title page to the document.
        
        Args:
            title: Document title
        """
        # Add title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Add date
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_text = f"Generated on: {datetime.now().strftime('%B %d, %Y')}"
        date_run = date_para.add_run(date_text)
        date_run.font.size = Pt(12)
        date_run.italic = True
        
        # Add divider
        divider_para = self.doc.add_paragraph()
        divider_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        divider_run = divider_para.add_run("________________________________________")
        divider_run.font.size = Pt(16)
        
        # Add description
        desc_para = self.doc.add_paragraph()
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_text = "IDMS Program Analysis"
        desc_run = desc_para.add_run(desc_text)
        desc_run.font.size = Pt(12)
    
    def add_table_of_contents(self) -> None:
        """Add a table of contents to the document"""
        toc_para = self.doc.add_paragraph()
        toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_run = toc_para.add_run("Table of Contents")
        toc_run.bold = True
        toc_run.font.size = Pt(16)
        
        # Add placeholder paragraph
        toc_para = self.doc.add_paragraph()
        
        # Insert TOC field
        run = toc_para.add_run()
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'begin')
        run._element.append(fld_char)
        
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._element.append(instr_text)
        
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'separate')
        run._element.append(fld_char)
        
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'end')
        run._element.append(fld_char)
        
        # Add page break after TOC
        self.doc.add_page_break()
    
    def convert_markdown_to_docx(self, markdown_content: str) -> None:
        """
        Convert markdown content to Word document elements.
        
        Args:
            markdown_content: Markdown content to convert
        """
        if not self.doc:
            raise ValueError("Document not initialized. Call create_document() first.")
        
        # Parse markdown and add to document
        lines = markdown_content.split('\n')
        i = 0
        
        # Track current list and code block state
        current_list = None
        in_code_block = False
        code_content = []
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Handle code blocks
            if line.startswith('```'):
                in_code_block = not in_code_block
                if not in_code_block and code_content:
                    # Add collected code content as a code block
                    code_para = self.doc.add_paragraph(style='Code')
                    code_text = '\n'.join(code_content)
                    code_para.add_run(code_text)
                    code_content = []
                i += 1
                continue
            
            if in_code_block:
                code_content.append(line)
                i += 1
                continue
            
            # Handle headings
            if line.startswith('# '):
                self.doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                self.doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                self.doc.add_heading(line[4:], 3)
            elif line.startswith('#### '):
                self.doc.add_heading(line[5:], 4)
            
            # Handle lists
            elif line.startswith('- ') or line.startswith('* '):
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:])
            elif re.match(r'^\d+\.\s', line):
                match = re.match(r'^\d+\.\s(.*)', line)
                if match:
                    p = self.doc.add_paragraph(style='List Number')
                    p.add_run(match.group(1))
            
            # Handle bold text
            elif line.startswith('**') and line.endswith('**') and len(line) > 4:
                p = self.doc.add_paragraph()
                p.add_run(line[2:-2]).bold = True
            
            # Handle italics
            elif line.startswith('*') and line.endswith('*') and len(line) > 2:
                p = self.doc.add_paragraph()
                p.add_run(line[1:-1]).italic = True
            
            # Handle horizontal rule
            elif line == '---' or line == '***' or line == '___':
                p = self.doc.add_paragraph()
                p.add_run('_' * 50)
            
            # Handle empty line
            elif line.strip() == '':
                if i > 0 and lines[i-1].strip() != '':
                    self.doc.add_paragraph()
            
            # Handle normal paragraphs with inline formatting
            else:
                p = self.doc.add_paragraph()
                
                # Process inline formatting
                j = 0
                while j < len(line):
                    # Bold text
                    if line[j:j+2] == '**' and '**' in line[j+2:]:
                        end_idx = line.find('**', j+2)
                        if end_idx != -1:
                            p.add_run(line[j+2:end_idx]).bold = True
                            j = end_idx + 2
                            continue
                    
                    # Italic text
                    if line[j:j+1] == '*' and '*' in line[j+1:]:
                        end_idx = line.find('*', j+1)
                        if end_idx != -1:
                            p.add_run(line[j+1:end_idx]).italic = True
                            j = end_idx + 1
                            continue
                    
                    # Code inline
                    if line[j:j+1] == '`' and '`' in line[j+1:]:
                        end_idx = line.find('`', j+1)
                        if end_idx != -1:
                            code_run = p.add_run(line[j+1:end_idx])
                            code_run.font.name = 'Courier New'
                            j = end_idx + 1
                            continue
                    
                    # Regular text
                    if j < len(line):
                        # Find the next special character
                        next_special = min(
                            (line.find(ch, j) if line.find(ch, j) != -1 else len(line))
                            for ch in ['**', '*', '`']
                        )
                        p.add_run(line[j:next_special])
                        j = next_special
            
            i += 1
    
    def add_program_metadata(self, metadata: Dict[str, Any]) -> None:
        """
        Add program metadata section to the document.
        
        Args:
            metadata: Dictionary containing program metadata
        """
        if not self.doc:
            raise ValueError("Document not initialized. Call create_document() first.")
        
        # Add metadata section
        self.doc.add_heading("Program Information", 1)
        
        # Create table for metadata
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Attribute"
        header_cells[1].text = "Value"
        
        # Apply header formatting
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add metadata rows
        for key, value in metadata.items():
            row_cells = table.add_row().cells
            
            # Format the key name (convert snake_case to Title Case)
            formatted_key = ' '.join(word.capitalize() for word in key.split('_'))
            row_cells[0].text = formatted_key
            
            # Format the value based on its type
            if isinstance(value, list):
                if value:
                    row_cells[1].text = '\n'.join(f"• {item}" for item in value)
                else:
                    row_cells[1].text = "None"
            else:
                row_cells[1].text = str(value) if value is not None else "None"
    
    def save_to_stream(self) -> BytesIO:
        """
        Save the document to a BytesIO stream.
        
        Returns:
            BytesIO stream containing the document
        """
        if not self.doc:
            raise ValueError("Document not initialized. Call create_document() first.")
        
        bio = BytesIO()
        self.doc.save(bio)
        bio.seek(0)
        return bio
    
    def get_download_link(self, filename: str) -> str:
        """
        Generate a download link for the document.
        
        Args:
            filename: Filename for download
            
        Returns:
            HTML string containing the download link
        """
        if not self.doc:
            raise ValueError("Document not initialized. Call create_document() first.")
        
        bio = self.save_to_stream()
        b64 = base64.b64encode(bio.read()).decode()
        
        # Create a styled download button
        styled_link = f"""
        <style>
        .download-button {{
            display: inline-block;
            padding: 0.7em 1.4em;
            margin: 0 0.3em 0.3em 0;
            border-radius: 0.3em;
            box-sizing: border-box;
            text-decoration: none;
            font-family: 'Roboto', sans-serif;
            font-weight: 500;
            color: #FFFFFF;
            background-color: #4CAF50;
            box-shadow: inset 0 -0.6em 1em -0.35em rgba(0,0,0,0.17),
                        inset 0 0.6em 2em -0.3em rgba(255,255,255,0.15),
                        inset 0 0 0em 0.05em rgba(255,255,255,0.12);
            text-align: center;
            position: relative;
        }}
        .download-button:hover {{
            background-color: #45a049;
        }}
        .download-button:active {{
            box-shadow: inset 0 0.6em 2em -0.3em rgba(0,0,0,0.15),
                        inset 0 0 0em 0.05em rgba(255,255,255,0.12);
        }}
        </style>
        <a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" 
           download="{filename}" class="download-button">
           📄 Download {filename}
        </a>
        """
        
        return styled_link


def markdown_to_docx(markdown_content: str, title: str, metadata: Optional[Dict[str, Any]] = None) -> BytesIO:
    """
    Convert markdown content to a Word document.
    
    Args:
        markdown_content: Markdown content to convert
        title: Document title
        metadata: Optional dictionary containing program metadata
        
    Returns:
        BytesIO stream containing the document
    """
    generator = DocxGenerator()
    generator.create_document(title)
    
    # Add metadata if provided
    if metadata:
        generator.add_program_metadata(metadata)
    
    # Add table of contents
    generator.add_table_of_contents()
    
    # Convert markdown content
    generator.convert_markdown_to_docx(markdown_content)
    
    return generator.save_to_stream()

def get_download_link_for_docx(docx_stream: BytesIO, filename: str) -> str:
    """
    Generate a download link for a Word document.
    
    Args:
        docx_stream: BytesIO stream containing the document
        filename: Filename for download
        
    Returns:
        HTML string containing the download link
    """
    b64 = base64.b64encode(docx_stream.read()).decode()
    
    # Create a styled download button
    styled_link = f"""
    <style>
    .download-button {{
        display: inline-block;
        padding: 0.7em 1.4em;
        margin: 0 0.3em 0.3em 0;
        border-radius: 0.3em;
        box-sizing: border-box;
        text-decoration: none;
        font-family: 'Roboto', sans-serif;
        font-weight: 500;
        color: #FFFFFF;
        background-color: #4CAF50;
        box-shadow: inset 0 -0.6em 1em -0.35em rgba(0,0,0,0.17),
                    inset 0 0.6em 2em -0.3em rgba(255,255,255,0.15),
                    inset 0 0 0em 0.05em rgba(255,255,255,0.12);
        text-align: center;
        position: relative;
    }}
    .download-button:hover {{
        background-color: #45a049;
    }}
    .download-button:active {{
        box-shadow: inset 0 0.6em 2em -0.3em rgba(0,0,0,0.15),
                    inset 0 0 0em 0.05em rgba(255,255,255,0.12);
    }}
    </style>
    <a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" 
       download="{filename}" class="download-button">
       📄 Download {filename}
    </a>
    """
    
    return styled_link