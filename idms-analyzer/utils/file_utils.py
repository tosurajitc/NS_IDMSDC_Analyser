"""
File utilities for handling file operations in the IDMS Analyzer application.
This module manages file uploads, validation, parsing, and export functionality.
"""

import os
import tempfile
import shutil
from pathlib import Path
import hashlib
import base64
from io import BytesIO
from typing import Optional, Dict, Any, List, Tuple, BinaryIO
import streamlit as st
from docx import Document
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Define constants
ALLOWED_EXTENSIONS = ['.txt', '.cbl', '.cob', '.cobol']
MAX_FILE_SIZE_MB = 10  # Maximum file size in MB

class FileProcessor:
    """Class for processing IDMS program files"""
    
    def __init__(self, cache_dir: Optional[str] = None):
        """Initialize with an optional cache directory"""
        self.cache_dir = cache_dir or os.path.join(tempfile.gettempdir(), "idms_analyzer_cache")
        os.makedirs(self.cache_dir, exist_ok=True)
    
    def validate_file(self, file: BinaryIO, filename: str) -> Tuple[bool, str]:
        """
        Validate if the uploaded file is acceptable.
        
        Args:
            file: File object
            filename: Name of the file
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file extension
        file_ext = Path(filename).suffix.lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            return False, f"Invalid file extension. Allowed extensions: {', '.join(ALLOWED_EXTENSIONS)}"
        
        # Check file size
        file.seek(0, os.SEEK_END)
        file_size_mb = file.tell() / (1024 * 1024)
        file.seek(0)
        
        if file_size_mb > MAX_FILE_SIZE_MB:
            return False, f"File too large. Maximum file size: {MAX_FILE_SIZE_MB}MB"
        
        # Basic content validation - check if it looks like COBOL/IDMS code
        content_sample = file.read(4096).decode('utf-8', errors='replace')
        file.seek(0)
        
        # Look for common COBOL/IDMS indicators in the content
        cobol_indicators = [
            "IDENTIFICATION DIVISION",
            "PROCEDURE DIVISION",
            "DATA DIVISION",
            "WORKING-STORAGE",
            "SCHEMA SECTION",
            "PROGRAM-ID"
        ]
        
        has_indicators = any(indicator in content_sample for indicator in cobol_indicators)
        if not has_indicators:
            return False, "File does not appear to be a valid COBOL/IDMS program"
        
        return True, ""
    
    def save_uploaded_file(self, file: BinaryIO, filename: str) -> str:
        """
        Save an uploaded file to the cache directory.
        
        Args:
            file: File object
            filename: Name of the file
            
        Returns:
            Path to the saved file
        """
        # Create a unique filename based on content hash and original name
        content = file.read()
        file.seek(0)
        
        content_hash = hashlib.md5(content).hexdigest()[:10]
        base_name = Path(filename).stem
        ext = Path(filename).suffix
        
        unique_filename = f"{base_name}_{content_hash}{ext}"
        file_path = os.path.join(self.cache_dir, unique_filename)
        
        with open(file_path, 'wb') as f:
            f.write(content)
        
        logger.info(f"Saved uploaded file to {file_path}")
        return file_path
    
    def read_file_content(self, file_path: str) -> str:
        """
        Read the content of a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            File content as a string
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            return content
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {str(e)}")
            raise
    
    def clear_cache(self, older_than_days: int = 1) -> None:
        """
        Clear cache files older than the specified number of days.
        
        Args:
            older_than_days: Number of days
        """
        import time
        
        current_time = time.time()
        count = 0
        
        for f in os.listdir(self.cache_dir):
            file_path = os.path.join(self.cache_dir, f)
            if os.path.isfile(file_path):
                file_age_days = (current_time - os.path.getmtime(file_path)) / (60 * 60 * 24)
                
                if file_age_days > older_than_days:
                    os.remove(file_path)
                    count += 1
        
        logger.info(f"Cleared {count} cache files older than {older_than_days} days")


def create_docx_from_markdown(markdown_content: str, title: str) -> Document:
    """
    Convert markdown content to a Word document.
    
    Args:
        markdown_content: Markdown content as a string
        title: Document title
        
    Returns:
        Document object
    """
    doc = Document()
    doc.add_heading(title, 0)
    
    # Very basic markdown parsing
    current_list = None
    in_code_block = False
    code_content = []
    
    lines = markdown_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            if not in_code_block and code_content:
                p = doc.add_paragraph()
                code_text = '\n'.join(code_content)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = 9
                code_content = []
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 4)
        
        # Handle lists
        elif line.startswith('- ') or line.startswith('* '):
            if current_list is None:
                current_list = doc.add_paragraph(style='List Bullet')
                current_list.add_run(line[2:])
            else:
                current_list.add_run('\n' + line[2:])
        elif line.startswith('1. ') or line.startswith('1) '):
            if current_list is None:
                current_list = doc.add_paragraph(style='List Number')
                current_list.add_run(line[3:])
            else:
                current_list.add_run('\n' + line[3:])
        
        # Handle text formatting
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            p.add_run(line.strip('*')).bold = True
        elif line.startswith('*') and line.endswith('*'):
            p = doc.add_paragraph()
            p.add_run(line.strip('*')).italic = True
        
        # Handle empty lines - reset list
        elif line.strip() == '':
            current_list = None
            doc.add_paragraph()
        
        # Handle normal paragraphs
        else:
            current_list = None
            doc.add_paragraph(line)
        
        i += 1
    
    return doc

def get_docx_download_link(doc: Document, filename: str) -> str:
    """
    Generate a link to download the document.
    
    Args:
        doc: Document object
        filename: Filename for download
        
    Returns:
        HTML link for downloading the document
    """
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    b64 = base64.b64encode(bio.read()).decode()
    download_link = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" class="download-button">Download {filename}</a>'
    
    # Add some basic styling to the download link
    styled_link = f"""
    <style>
    .download-button {{
        display: inline-block;
        padding: 0.5em 1em;
        background-color: #4CAF50;
        color: white;
        text-decoration: none;
        border-radius: 4px;
        font-weight: bold;
        margin: 10px 0;
    }}
    .download-button:hover {{
        background-color: #45a049;
    }}
    </style>
    {download_link}
    """
    
    return styled_link

def detect_idms_program_type(file_content: str) -> str:
    """
    Detect if an IDMS program is DB-only or DC (online) based on its content.
    
    Args:
        file_content: The IDMS program code as a string
        
    Returns:
        "IDMS-DC" if it's an IDMS-DC program, "IDMS-DB" if it's a DB-only program
    """
    # Look for indicators of DC programs
    dc_indicators = [
        "PROTOCOL.  MODE IS IDMS-DC",
        "MAP SECTION",
        "MAP OUT",
        "MAP IN",
        "MODIFY MAP",
        "DC RETURN",
        "TRANSFER CONTROL",
        "DC-AID-CONDITION-NAMES"
    ]
    
    for indicator in dc_indicators:
        if indicator in file_content:
            return "IDMS-DC"
    
    # Even if no explicit DC indicators, check for schema
    if "SCHEMA SECTION" in file_content:
        return "IDMS-DB"
    
    # Default to IDMS-DB if no clear indicators
    return "IDMS-DB"

def extract_program_information(file_content: str) -> Dict[str, Any]:
    """
    Extract basic information from IDMS program.
    
    Args:
        file_content: The IDMS program code as a string
        
    Returns:
        Dictionary with program information
    """
    info = {
        "program_id": None,
        "program_type": detect_idms_program_type(file_content),
        "remarks": [],
        "history": [],
        "line_count": len(file_content.split('\n'))
    }
    
    # Extract program ID
    program_id_lines = [line for line in file_content.split('\n') if "PROGRAM-ID." in line]
    if program_id_lines:
        parts = program_id_lines[0].split('.')
        if len(parts) > 1:
            info["program_id"] = parts[1].strip()
    
    # Extract remarks
    in_remarks = False
    for line in file_content.split('\n'):
        line = line.strip()
        
        if "*REMARKS." in line:
            in_remarks = True
            continue
        
        if in_remarks and line.startswith('*'):
            if "*HISTORY:" in line:
                in_remarks = False
                continue
            
            remark = line[1:].strip()
            if remark and not remark.startswith('--'):
                info["remarks"].append(remark)
        elif in_remarks:
            in_remarks = False
    
    # Extract history
    in_history = False
    for line in file_content.split('\n'):
        line = line.strip()
        
        if "*HISTORY:" in line:
            in_history = True
            continue
        
        if in_history and line.startswith('*'):
            history_item = line[1:].strip()
            if history_item and not history_item.startswith('--'):
                info["history"].append(history_item)
        elif in_history and not line.startswith('*'):
            in_history = False
    
    return info