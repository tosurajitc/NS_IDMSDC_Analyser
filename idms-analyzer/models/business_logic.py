from typing import List, Optional, Dict, Any
from pydantic import BaseModel, Field, validator
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class IntegrationPoint(BaseModel):
    """Model representing an integration point in IDMS programs."""
    name: str = Field(..., description="Name of the integration point")
    description: str = Field(..., description="Description of the integration point")
    type: str = Field(..., description="Type of integration (e.g., 'Database', 'API', 'File')")
    direction: str = Field(..., description="Direction of data flow (e.g., 'Input', 'Output', 'Both')")
    
    @validator('direction')
    def validate_direction(cls, v):
        valid_directions = ['Input', 'Output', 'Both']
        if v not in valid_directions:
            raise ValueError(f"Direction must be one of {valid_directions}")
        return v

class Validation(BaseModel):
    """Model representing a validation rule in the program."""
    field: str = Field(..., description="Field or data element being validated")
    rule: str = Field(..., description="Description of the validation rule")
    error_handling: Optional[str] = Field(None, description="How errors are handled")

class SpecialCase(BaseModel):
    """Model representing a special case or exception in program logic."""
    condition: str = Field(..., description="Condition that triggers the special case")
    handling: str = Field(..., description="How the special case is handled")
    notes: Optional[str] = Field(None, description="Additional notes about the special case")

class BusinessRule(BaseModel):
    """Model representing a core business rule in the program."""
    rule_id: Optional[str] = Field(None, description="Identifier for the rule")
    description: str = Field(..., description="Description of the business rule")
    implementation: str = Field(..., description="How the rule is implemented in code")
    
    @validator('rule_id', pre=True, always=True)
    def set_rule_id(cls, v, values):
        if v:
            return v
        # Generate a rule ID based on description if not provided
        if 'description' in values:
            words = re.sub(r'[^\w\s]', '', values['description']).split()[:3]
            return 'RULE_' + '_'.join(words).upper()
        return 'RULE_UNKNOWN'

class BusinessLogic(BaseModel):
    """Main model for storing extracted business logic from IDMS programs."""
    program_name: str = Field(..., description="Name of the IDMS program")
    program_type: str = Field(..., description="Type of program (DC/Online or DB/Batch)")
    program_purpose: str = Field(..., description="Overall purpose of the program")
    core_rules: List[BusinessRule] = Field(default_factory=list, description="Core business rules")
    validations: List[Validation] = Field(default_factory=list, description="Validation rules")
    special_cases: List[SpecialCase] = Field(default_factory=list, description="Special cases and exceptions")
    integration_points: List[IntegrationPoint] = Field(default_factory=list, description="Integration points")
    additional_notes: Optional[str] = Field(None, description="Any additional notes or comments")
    
    @validator('program_type')
    def validate_program_type(cls, v):
        valid_types = ['DC', 'DB', 'Online', 'Batch']
        if v not in valid_types:
            raise ValueError(f"Program type must be one of {valid_types}")
        return v
    
    def generate_word_document(self) -> Document:
        """
        Generate a comprehensive Word document for the business logic.
        
        Returns:
            Document: Fully formatted Word document
        """
        # Create a new Document
        doc = Document()
        
        # Title Page
        title_page = doc.add_paragraph()
        title_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_page.add_run(f"{self.program_name} Business Logic")
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Add page break
        doc.add_page_break()
        
        # Program Information Section
        doc.add_heading("Program Information", level=1)
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        
        # Populate program information
        info_rows = [
            ("Program Name", self.program_name),
            ("Program Type", self.program_type),
            ("Program Purpose", self.program_purpose)
        ]
        
        for i, (label, value) in enumerate(info_rows):
            cells = info_table.rows[i].cells
            cells[0].text = label
            cells[1].text = value
        
        # Core Business Rules Section
        doc.add_heading("Core Business Rules", level=1)
        if not self.core_rules:
            doc.add_paragraph("No core business rules identified.")
        else:
            for rule in self.core_rules:
                doc.add_heading(rule.rule_id, level=2)
                doc.add_paragraph(f"Description: {rule.description}")
                doc.add_paragraph(f"Implementation: {rule.implementation}")
        
        # Validations Section
        doc.add_heading("Data Validation Rules", level=1)
        if not self.validations:
            doc.add_paragraph("No validation rules identified.")
        else:
            validation_table = doc.add_table(rows=1, cols=3)
            validation_table.style = 'Table Grid'
            
            # Table headers
            headers = validation_table.rows[0].cells
            headers[0].text = "Field"
            headers[1].text = "Validation Rule"
            headers[2].text = "Error Handling"
            
            # Add validation rows
            for validation in self.validations:
                row_cells = validation_table.add_row().cells
                row_cells[0].text = validation.field
                row_cells[1].text = validation.rule
                row_cells[2].text = validation.error_handling or "N/A"
        
        # Special Cases Section
        doc.add_heading("Special Processing Rules", level=1)
        if not self.special_cases:
            doc.add_paragraph("No special cases identified.")
        else:
            for case in self.special_cases:
                doc.add_heading(case.condition, level=2)
                doc.add_paragraph(f"Handling: {case.handling}")
                if case.notes:
                    doc.add_paragraph(f"Notes: {case.notes}")
        
        # Integration Points Section
        doc.add_heading("Integration Points", level=1)
        if not self.integration_points:
            doc.add_paragraph("No integration points identified.")
        else:
            for point in self.integration_points:
                doc.add_heading(point.name, level=2)
                point_table = doc.add_table(rows=3, cols=2)
                point_table.style = 'Table Grid'
                
                point_rows = [
                    ("Type", point.type),
                    ("Direction", point.direction),
                    ("Description", point.description)
                ]
                
                for i, (label, value) in enumerate(point_rows):
                    cells = point_table.rows[i].cells
                    cells[0].text = label
                    cells[1].text = value
        
        # Additional Notes Section
        if self.additional_notes:
            doc.add_heading("Additional Notes", level=1)
            doc.add_paragraph(self.additional_notes)
        
        return doc

class TestCase(BaseModel):
    """
    Represents a single test scenario derived from business logic
    """
    test_id: str = Field(..., description="Unique test case identifier")
    title: str = Field(..., description="Title of the test case")
    description: str = Field(..., description="Detailed description of the test scenario")
    prerequisites: List[str] = Field(default_factory=list, description="Conditions that must be met before test execution")
    test_data: Dict[str, Any] = Field(default_factory=dict, description="Input data for the test case")
    steps: List[str] = Field(..., description="Sequence of steps to execute the test")
    expected_results: List[str] = Field(..., description="Expected outcomes of the test")
    related_rules: List[str] = Field(default_factory=list, description="Business rules related to this test case")

    def generate_word_section(self, doc: Document):
        """
        Generate a section in the Word document for this test case
        
        Args:
            doc: The Word document to add the test case to
        """
        # Add page break if not the first section
        if doc.paragraphs:
            doc.add_page_break()
        
        # Test Case Header
        doc.add_heading(f"Test Case: {self.test_id}", level=2)
        doc.add_heading(self.title, level=3)
        
        # Description
        doc.add_paragraph(f"Description: {self.description}")
        
        # Prerequisites
        doc.add_heading("Prerequisites", level=4)
        if not self.prerequisites:
            doc.add_paragraph("None")
        else:
            for prereq in self.prerequisites:
                doc.add_paragraph(prereq, style='List Bullet')
        
        # Test Data
        doc.add_heading("Test Data", level=4)
        if not self.test_data:
            doc.add_paragraph("None")
        else:
            for key, value in self.test_data.items():
                doc.add_paragraph(f"{key}: {value}")
        
        # Steps
        doc.add_heading("Steps", level=4)
        for i, step in enumerate(self.steps, 1):
            doc.add_paragraph(f"{i}. {step}")
        
        # Expected Results
        doc.add_heading("Expected Results", level=4)
        for result in self.expected_results:
            doc.add_paragraph(result, style='List Bullet')
        
        # Related Rules
        if self.related_rules:
            doc.add_heading("Related Business Rules", level=4)
            doc.add_paragraph(", ".join(self.related_rules))

class TestScript(BaseModel):
    """
    Represents a collection of test cases for an IDMS program
    """
    program_name: str
    test_cases: List[TestCase] = Field(default_factory=list)

    def generate_word_document(self) -> Document:
        """
        Generate a comprehensive Word document for the test script
        
        Returns:
            Document: Fully formatted Word document with test cases
        """
        # Create a new Document
        doc = Document()
        
        # Title Page
        title_page = doc.add_paragraph()
        title_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_page.add_run(f"Test Script for {self.program_name}")
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Add a page break
        doc.add_page_break()
        
        # Add test script overview
        doc.add_heading("Test Script Overview", level=1)
        doc.add_paragraph(f"Program: {self.program_name}")
        doc.add_paragraph(f"Total Test Cases: {len(self.test_cases)}")
        
        # Add each test case
        for test_case in self.test_cases:
            test_case.generate_word_section(doc)
        
        return doc